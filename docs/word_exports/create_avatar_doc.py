import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set document properties
doc.core_properties.title = "EdPsych AI Avatar Video Scripts and Generation Guide"
doc.core_properties.author = "EdPsych Connect"
doc.core_properties.comments = "Complete guide for generating AI avatar videos"

# Add title
title = doc.add_heading('EdPsych AI Avatar Video Scripts and Generation Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add introduction
doc.add_paragraph('This document contains all scripts and instructions for generating AI avatar videos for the EdPsych AI Education Platform. It includes detailed scripts for each avatar type, generation instructions, and recommended platforms.')

# Add table of contents header
toc_heading = doc.add_heading('Table of Contents', 1)
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Placeholder for TOC (will be manually updated in Word)
doc.add_paragraph('Right-click here and select "Update Field" to update the table of contents.')

# Add section: Implementation Guide
doc.add_heading('Implementation Guide', 1)
doc.add_paragraph('The following guide provides step-by-step instructions for generating and implementing AI avatar videos for the EdPsych AI Education Platform.')

# Read and add implementation guide content
guide_path = '/home/ubuntu/edpsych_work/EdPsych-AI-Education-Platform/docs/ai_avatar_assets/COMPLETE_IMPLEMENTATION_GUIDE.md'
with open(guide_path, 'r') as f:
    guide_content = f.read()
    
# Process markdown headings and content
lines = guide_content.split('\n')
current_section = ""
for line in lines:
    if line.startswith('# '):
        # Skip the title as we've already added it
        continue
    elif line.startswith('## '):
        # Level 2 heading
        doc.add_heading(line[3:], 2)
    elif line.startswith('### '):
        # Level 3 heading
        doc.add_heading(line[4:], 3)
    elif line.startswith('```'):
        # Code block - we'll just add as regular text for now
        current_section = "code" if current_section != "code" else ""
    elif line.startswith('- '):
        # Bullet point
        doc.add_paragraph(line[2:], style='List Bullet')
    elif line.startswith('1. '):
        # Numbered list
        doc.add_paragraph(line[3:], style='List Number')
    else:
        # Regular paragraph
        if line.strip() and current_section != "code":
            doc.add_paragraph(line)
        elif current_section == "code" and line.strip():
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

# Add section: Recommended Platforms
doc.add_heading('Recommended Platforms for AI Avatar Generation', 1)
doc.add_paragraph('Based on our evaluation, the following platforms are recommended for generating AI avatar videos:')

# Add platform details
platforms = [
    {
        'name': 'HeyGen',
        'url': 'https://www.heygen.com/',
        'description': 'Primary recommended platform with excellent avatar quality and voice cloning capabilities.',
        'strengths': [
            'High-quality realistic avatars',
            'Excellent voice cloning',
            'Good API integration',
            'Commercial usage rights available',
            'Supports multiple languages'
        ],
        'pricing': 'Subscription-based with various tiers. Business plans start at approximately £79/month.'
    },
    {
        'name': 'D-ID',
        'url': 'https://www.d-id.com/',
        'description': 'Secondary recommended platform with good avatar quality and integration options.',
        'strengths': [
            'Good avatar quality',
            'Strong API capabilities',
            'Multiple background options',
            'Supports various languages',
            'Good for batch processing'
        ],
        'pricing': 'Pay-as-you-go and subscription options available. Business plans start at approximately £49/month.'
    },
    {
        'name': 'Synthesia',
        'url': 'https://www.synthesia.io/',
        'description': 'Tertiary option with good quality and extensive template library.',
        'strengths': [
            'Large template library',
            'Good for educational content',
            'Multiple avatar options',
            'User-friendly interface',
            'Good customer support'
        ],
        'pricing': 'Subscription-based. Business plans start at approximately £65/month.'
    }
]

for platform in platforms:
    doc.add_heading(platform['name'], 2)
    doc.add_paragraph(f"Website: {platform['url']}")
    doc.add_paragraph(platform['description'])
    
    doc.add_heading('Strengths:', 3)
    for strength in platform['strengths']:
        doc.add_paragraph(strength, style='List Bullet')
    
    doc.add_paragraph(f"Pricing: {platform['pricing']}")

# Add section: Avatar Scripts
doc.add_heading('Avatar Scripts', 1)
doc.add_paragraph('The following sections contain all scripts for the different avatar types and video clips.')

# Function to add script content from file
def add_script_content(doc, file_path, section_title):
    doc.add_heading(section_title, 2)
    
    try:
        with open(file_path, 'r') as f:
            content = f.read()
            
        lines = content.split('\n')
        in_script_section = False
        current_heading = ""
        
        for line in lines:
            if line.startswith('# '):
                # Main title - skip as we've added our own
                continue
            elif line.startswith('## '):
                # Level 2 heading
                current_heading = line[3:]
                doc.add_heading(current_heading, 3)
                in_script_section = True
            elif line.startswith('### '):
                # Level 3 heading
                doc.add_heading(line[4:], 4)
            elif line.startswith('[*') and line.endswith('*]'):
                # Production note
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.italic = True
                run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
            elif in_script_section and line.strip():
                # Regular paragraph in script section
                doc.add_paragraph(line)
    except FileNotFoundError:
        doc.add_paragraph(f"Script file not found: {file_path}")

# Add Teacher Avatar Scripts
add_script_content(
    doc, 
    '/home/ubuntu/edpsych_work/EdPsych-AI-Education-Platform/docs/ai_avatar_assets/teacher_avatar_clips.md',
    'Teacher Avatar Scripts (Dr. Scott)'
)

# Add Student Peer Avatar Scripts
add_script_content(
    doc, 
    '/home/ubuntu/edpsych_work/EdPsych-AI-Education-Platform/docs/ai_avatar_assets/student_peer_avatar_clips.md',
    'Student Peer Avatar Scripts'
)

# Add Special Needs Support Avatar Scripts
add_script_content(
    doc, 
    '/home/ubuntu/edpsych_work/EdPsych-AI-Education-Platform/docs/ai_avatar_assets/special_needs_support_avatar_clips.md',
    'Special Needs Support Avatar Scripts'
)

# Add Subject Specialist Avatar Scripts
add_script_content(
    doc, 
    '/home/ubuntu/edpsych_work/EdPsych-AI-Education-Platform/docs/ai_avatar_assets/subject_specialist_avatar_clips.md',
    'Subject Specialist Avatar Scripts'
)

# Add Full-Length Video Scripts
doc.add_heading('Full-Length Video Scripts', 2)

# Executive Summary Script
add_script_content(
    doc, 
    '/home/ubuntu/edpsych_work/EdPsych-AI-Education-Platform/docs/video_scripts/executive_summary_script.md',
    'Executive Summary Video Script'
)

# Platform Features Overview Script
add_script_content(
    doc, 
    '/home/ubuntu/edpsych_work/EdPsych-AI-Education-Platform/docs/video_scripts/platform_features_overview_script.md',
    'Platform Features Overview Video Script'
)

# Educational Psychology Foundations Script
add_script_content(
    doc, 
    '/home/ubuntu/edpsych_work/EdPsych-AI-Education-Platform/docs/video_scripts/educational_psychology_foundations_script.md',
    'Educational Psychology Foundations Video Script'
)

# Add appendix with file structure
doc.add_heading('Appendix: File Structure for Avatar Videos', 1)
doc.add_paragraph('The following file structure should be used for organizing the generated avatar videos:')

file_structure = """
/public/videos/avatars/
├── teacher/         # Videos featuring Dr. Scott Ighavongbe-Patrick
│   ├── introduction/
│   ├── instructional/
│   ├── emotional_support/
│   ├── transition/
│   └── full_length/
├── student/         # Videos featuring student peer avatars
│   ├── introduction/
│   ├── instructional/
│   ├── emotional_support/
│   ├── transition/
│   └── peer_explanation/
├── special_needs/   # Videos featuring special needs support avatars
│   ├── introduction/
│   ├── instructional/
│   ├── emotional_support/
│   ├── transition/
│   └── specialized_support/
└── specialist/      # Videos featuring subject specialist avatars
    ├── introduction/
    ├── instructional/
    ├── subject_specific/
    ├── advanced_topic/
    ├── interdisciplinary/
    └── curriculum_connection/
"""

p = doc.add_paragraph()
run = p.add_run(file_structure)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# Add naming convention
doc.add_heading('File Naming Convention', 2)
doc.add_paragraph('Files should follow this naming convention:')
p = doc.add_paragraph()
run = p.add_run('[avatar_type]_[avatar_name]_[category]_[clip_id].mp4')
run.font.name = 'Courier New'

doc.add_paragraph('Examples:')
examples = [
    'teacher_dr_scott_introduction_welcome.mp4',
    'student_alex_instructional_assessment_tools.mp4',
    'special_needs_sam_emotional_support_managing_frustration.mp4',
    'specialist_dr_maya_subject_specific_understanding_fractions.mp4'
]

for example in examples:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(example)
    run.font.name = 'Courier New'

# Save the document
output_path = '/home/ubuntu/edpsych_work/EdPsych-AI-Education-Platform/docs/word_exports/EdPsych_AI_Avatar_Scripts_and_Guide.docx'
doc.save(output_path)

print(f"Document created successfully at: {output_path}")
